#!/usr/bin/python

import os
import random
import datetime
import random
import time
import json
import getpass
from pathlib import Path

import paramiko
import xlwt
import xml.etree.cElementTree as ET
import sqlite3
from cryptography.fernet import Fernet
from docx import Document
from fpdf import FPDF
from etc.api.googledrv.gdrive_folder import *
from etc.serverx.config.config_server import *
from etc.aws.bucket_s3 import *


# Global variables/Variáveis globais.
path_atual_ct = str(pathlib.Path(__file__).parent.absolute())
path_ct_final = path_atual_ct.replace('/etc','')
Path(f'/home/{getpass.getuser()}/Documents/Oblivion').mkdir(parents=True, exist_ok=True)


def ssh_enviar_arquivo(arquivo, destino, key=None):
    """
    Sends file by SSH/Envia arquivo por SSH.

    :param arquivo: File/Arquivo.
    :param destino: Destiny/Destino.
    """
    with open(f'{path_ct_final}/etc/serverx/config/ssh_hosts.txt', 'r') as hosts_ssh:
        final_ssh = []
        conteudo_hosts = hosts_ssh.read()
        conteudo_hosts = conteudo_hosts.split('\n')

        for e in conteudo_hosts:

            if 'amazonaws' in e:

                try:
                    e = e.split(';')

                    if e[5] != key:
                        break

                    # If you want that the Oblivion creates the folder, uncommnet this:
                    # Se vc quiser que o Oblivion crie uma pasta, descomente isso:
                    #
                    # stdin, stdout, stderr = ssh.exec_command("mkdir leak")

                    ssh = paramiko.SSHClient()
                    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
                    ssh.connect(str(e[0]),
                                username=str(e[2]),
                                key_filename=f'''{path_ct_final}/etc/serverx/pems/{e[6]}''')
                    sftp = ssh.open_sftp()
                    sftp.put(arquivo, f'{e[4]}/{destino}')
                    sftp.close()
                    ssh.close()

                except:
                    pass

            else:
                try:
                    e = e.split(';')

                    if e[5] != key:
                        break

                    s = paramiko.SSHClient()
                    s.set_missing_host_key_policy(paramiko.AutoAddPolicy())
                    s.connect(e[0], int(e[1]), e[2], e[3])
                    sftp = s.open_sftp()
                    sftp.put(arquivo, f'{e[4]}/{destino}')
                    sftp.close()
                    s.close()

                except:
                    pass


def ocultar_senhas(senhas):
    """
    Ocults the last characters of password/Oculta os últimos caractéres da senha.

    :param senhas: Password/Senha.
    """
    if senhas != 'None':
        #senhas = senhas.encode('utf-16', 'replace').decode('utf-16')
        tamanho = int(len(senhas)/2)
        tamanho = -tamanho
        repor = senhas[tamanho:]
        senhas = senhas.replace(repor,'*****')
        return senhas

    return senhas


def criptografar(arquivo):
    """
    Crypt a file/Criptografa um arquivo.

    :param arquivo: Path of file/Local do arquivo.
    """
    #key = b'mrYXrs2rnYanAx88wDdqvgFc6aapqU5a-CnHSp_AAmk='
    with open(f'{path_ct_final}/etc/key_crypt.txt', 'r') as pegar_key:
        key = pegar_key.read()

    input_file = arquivo
    output_file = arquivo + '.encrypted'

    with open(input_file, 'rb') as f:
        data = f.read()

    fernet = Fernet(key)
    encrypted = fernet.encrypt(data)

    with open(output_file, 'wb') as f:
        f.write(encrypted)


def descriptografar(arquivo):
    """
    Decrypt a file/Descriptografa um arquivo.

    :param arquivo: Path of file/Local do arquivo.
    """
    key = b'mrYXrs2rnYanAx88wDdqvgFc6aapqU5a-CnHSp_AAmk='
    input_file = arquivo + '.encrypted'
    output_file = arquivo

    with open(input_file, 'rb') as f:
        data = f.read()

    fernet = Fernet(key)

    try:
        decrypted = fernet.decrypt(data)

        with open(output_file, 'wb') as f:
            f.write(decrypted)

    except:
        pass


def string_grande(criar=None, criar_ocult=None, criar_cript=None, documentos=None, nome=None, data_atual=None,
                  binario=None, ssh_enviar=None, id_pasta=None, gdrive=None, aws_s3_gerar=None, serverx=None,
                  key=None):
    """
    Creates a txt file with the string if the string is big too much (> len 500 aprox)/Cria uma arquivo txt caso a string em questão
    for muito grande (> len 500 aprox).
    """
    if criar == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if criar_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', 'w+', encoding='utf-16') as arquivo:
                #binario = binario.encode('utf-16', 'replace').decode('utf-16')
                arquivo.write(f'{binario}\n')

        if criar_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', 'w+', encoding='utf-16') as arquivo:
                #binario = binario.encode('utf-16', 'replace').decode('utf-16')
                arquivo.write(f'{binario}\n')

        if criar_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted', 'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:
            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', 'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)


def gerar_resultados(
        serverx=False, key=None,
        ssh_enviar=False, aws_s3_gerar=False,
        nome=None, gdrive=False, conteudo=None,
        txt_f=False, txt_cript=False, txt_ocult=False,
        docx_f=False, docx_cript=False, docx_ocult=False,
        pdf_f=False, pdf_cript=False, pdf_ocult=False,
        xlsx_f=False, xlsx_cript=False, xlsx_ocult=False,
        json_f=False, json_cript=False, json_ocult=False,
        html_f=False, html_cript=False, html_ocult=False,
        xsl_f=False, xsl_cript=False, xsl_ocult=False,
        db_f=False, db_cript=False, db_ocult=False,
        gd_txt_f=False, gd_txt_cript=False, gd_txt_ocult=False,
        gd_docx_f=False, gd_docx_cript=False, gd_docx_ocult=False,
        gd_pdf_f=False, gd_pdf_cript=False, gd_pdf_ocult=False,
        gd_xlsx_f=False, gd_xlsx_cript=False, gd_xlsx_ocult=False,
        gd_json_f=False, gd_json_cript=False, gd_json_ocult=False,
        gd_html_f=False, gd_html_cript=False, gd_html_ocult=False,
        gd_xsl_f=False, gd_xsl_cript=False, gd_xsl_ocult=False,
        gd_db_f=False, gd_db_cript=False, gd_db_ocult=False
):
    """
    Generates final file/Gerar arquivo final.

    :param serverx: If is Oblivion Server or Oblivion Client/Se é o Cliente Oblivion ou o Server Oblivion.
    :param ssh_gerar: Sends file by SSH/Enviar arquivo por SSH.
    :param aws_s3_gerar: Sends file to S3 Bucket of AWS/Enviar arquivo para o Bucket S3 da AWS.
    :param nome: Name of file/Nome do arquivo.
    :param gdr: Uploads to Google Drive/Sobe arquivo para o Google Drive.
    :param conteudo: Data/Dados.

    :param ..._f: Generates file/Gera arquivo.
    :param ..._cript: Generates encrypted file/Gera um arquivo criptografado.
    :param ..._ocult: Generates a file with oculted passwords/Gera um arquivo com a senhas ocultas.

    Examples:

    ..._f + ..._cript = Creates encrypted file/Cria arquivo criptografado.
    ..._f + ..._cript + ..._ocult = Creates encrypted file with oculted password/Cria arquivo criptografado com as senhas
    ocultas.
    """

    data_atual = datetime.date.today()
    definir = os.environ

    with open(f'{path_ct_final}/etc/api/googledrv/id_folder.txt', 'r') as pegar_id_pasta:
        id_pasta = str(pegar_id_pasta.read())

    for valor, item in definir.items():

        if valor == 'APPDATA':
            lista_item = item.split('\\')
            usuario = lista_item[2]
    
    usuario = getpass.getuser()
    documentos = f'/home/{usuario}/Documents'

    if txt_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if txt_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', 'w+', encoding='utf-16') as arquivo:
                for binario in conteudo:

                    if not 'N/A' in binario:
                        #binario = binario.encode('utf-16', 'replace').decode('utf-16')
                        binario_repor = binario.split(':')
                        senha = binario_repor[1]
                        senha_ocult = ocultar_senhas(senha)
                        binario = binario.replace(senha, senha_ocult)
                        arquivo.write(f'{binario}\n')

        if txt_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', 'w+', encoding='utf-16') as arquivo:
                for binario in conteudo:
                    #binario = binario.encode('utf-16', 'replace').decode('utf-16')
                    arquivo.write(f'{binario}\n')

        if txt_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted', 'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:
            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', 'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.txt',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.txt', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)


    if docx_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if docx_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            document = Document()
            for binario in conteudo:

                if not 'N/A' in binario:
                    binario = binario.encode('utf-16', 'replace').decode('utf-16')
                    binario_repor = binario.split(':')
                    senha = binario_repor[1]
                    senha_ocult = ocultar_senhas(senha)
                    binario = binario.replace(senha, senha_ocult)
                    document.add_paragraph(binario)
            document.save(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx')

        if docx_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            document = Document()
            for binario in conteudo:
                binario = binario.encode('utf-16', 'replace').decode('utf-16')
                document.add_paragraph(binario)
            document.save(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx')

        if docx_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.docx.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.docx.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.docx',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.docx',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.docx', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

    if pdf_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if pdf_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            dummytext = ""
            for binario in conteudo:

                if not 'N/A' in binario:
                    binario = binario.encode('latin-1', 'replace').decode('latin-1')
                    binario_repor = binario.split(':')
                    senha = binario_repor[1]
                    senha_ocult = ocultar_senhas(senha)
                    binario = binario.replace(senha, senha_ocult)
                    dummytext += f'{binario}\n\n'

            pdf = FPDF()
            pdf.add_page()
            pdf.set_xy(5, 5)
            pdf.set_font('arial', 'B', 12.0)
            pdf.multi_cell(0, 5, dummytext)
            pdf.output(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf', 'F')

        if pdf_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            dummytext = ""
            for binario in conteudo:
                binario = binario.encode('latin-1', 'replace').decode('latin-1')
                dummytext += f'{binario}\n\n'

            pdf = FPDF()
            pdf.add_page()
            pdf.set_xy(5, 5)
            pdf.set_font('arial', 'B', 12.0)
            pdf.multi_cell(0, 5, dummytext)
            pdf.output(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf', 'F')

        if pdf_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.pdf', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

    if xlsx_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if xlsx_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            x = 0
            y = 0
            style0 = xlwt.easyxf('font: name Times New Roman, color-index red, bold on',
                                 num_format_str='#,##0.00')
            style1 = xlwt.easyxf(num_format_str='D-MMM-YY')
            wb = xlwt.Workbook()
            ws = wb.add_sheet('A Test Sheet')

            try:
                for binario in conteudo:
                    if not 'N/A' in binario:
                        binario = binario.encode('utf-16', 'replace').decode('utf-16')
                        binario_repor = binario.split(':')
                        senha = binario_repor[1]
                        senha_ocult = ocultar_senhas(senha)
                        binario = binario.replace(senha, senha_ocult)
                    binario_x = binario.split(':')

                    if y >= 3:
                        y = 0

                    for kls in binario_x:
                        ws.write(x, y, kls)
                        y += 1

                    x += 1
                    wb.save(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls')

            except:
                string_grande(criar=xlsx_f, criar_ocult=xlsx_ocult, criar_cript=xlsx_cript, documentos=documentos,
                              nome=nome, data_atual=data_atual, binario=conteudo, ssh_enviar=ssh_enviar,
                              id_pasta=id_pasta,
                              gdrive=gdrive, aws_s3_gerar=aws_s3_gerar, serverx=serverx, key=key)

        if xlsx_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            hash_file = random.getrandbits(23)
            x = 0
            y = 0
            style0 = xlwt.easyxf('font: name Times New Roman, color-index red, bold on',
                                 num_format_str='#,##0.00')
            style1 = xlwt.easyxf(num_format_str='D-MMM-YY')

            wb = xlwt.Workbook()
            ws = wb.add_sheet('A Test Sheet')

            try:
                for binario in conteudo:
                    binario = binario.encode('utf-16', 'replace').decode('utf-16')
                    binario_x = binario.split(':')

                    if y >= 3:
                        y = 0

                    for kls in binario_x:
                        ws.write(x, y, kls)
                        y += 1
                    x += 1

                    wb.save(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls')

            except:
                string_grande(criar=xlsx_f, criar_ocult=xlsx_ocult, criar_cript=xlsx_cript, documentos=documentos,
                              nome=nome, data_atual=data_atual, binario=conteudo, ssh_enviar=ssh_enviar,
                              id_pasta=id_pasta,
                              gdrive=gdrive, aws_s3_gerar=aws_s3_gerar, serverx=serverx, key=key)

        if xlsx_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xls.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xls.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xls',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xls',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xls', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

    if json_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if json_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json', 'w', encoding='utf-16') as arquivo:
                temp_lista = []
                for binario in conteudo:

                    if not 'N/A' in binario:
                        #binario = binario.encode('utf-16', 'replace').decode('utf-16')
                        binario_repor = binario.split(':')
                        senha = binario_repor[1]
                        senha_ocult = ocultar_senhas(senha)
                        binario = binario.replace(senha, senha_ocult)

                        tmp_lista = []
                        final_list = []
                        tmp_lista.append(binario)
                        temp_lista.append(tmp_lista)
                json.dump(temp_lista, arquivo)

        if json_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json', 'w', encoding='utf-16') as arquivo:
                temp_lista = []
                for binario in conteudo:

                    if not 'N/A' in binario:
                        #binario = binario.encode('utf-16', 'replace').decode('utf-16')
                        tmp_lista = []
                        final_list = []
                        tmp_lista.append(binario)
                        temp_lista.append(tmp_lista)
                json.dump(temp_lista, arquivo)

        if json_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.json.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.json.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.json',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.json',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.json', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

    if html_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if html_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html', 'w+', encoding='utf-16') as arquivo:
                arquivo.write(f'<!DOCTYPE html>')
                for binario in conteudo:

                    if not 'N/A' in binario:
                        binario = binario.encode('utf-16', 'replace').decode('utf-16')
                        binario_repor = binario.split(':')
                        senha = binario_repor[1]
                        senha_ocult = ocultar_senhas(senha)
                        binario = binario.replace(senha, senha_ocult)
                        arquivo.write(f'<div>{binario}</div>')
                        arquivo.write(f'</br>')

        if html_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            with open(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html', 'w+', encoding='utf-16') as arquivo:
                arquivo.write(f'<!DOCTYPE html>')
                for binario in conteudo:
                    binario = binario.encode('utf-16', 'replace').decode('utf-16')
                    arquivo.write(f'<div>{binario}</div>')
                    arquivo.write(f'</br>')

        if html_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.html.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.html.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.html',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.html',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.html', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)


    if xsl_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)


        if xsl_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            root = ET.Element("oblivion")
            doc = ET.SubElement(root, "leak")
            for binario in conteudo:

                if not 'N/A' in binario:
                    binario = binario.encode('utf-16', 'replace').decode('utf-16')
                    binario_repor = binario.split(':')
                    senha = binario_repor[1]
                    senha_ocult = ocultar_senhas(senha)
                    binario = binario.replace(senha, senha_ocult)
                    ET.SubElement(doc, f"field", name="leak").text = binario
            tree = ET.ElementTree(root)
            tree.write(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml')

        if xsl_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':
            root = ET.Element("oblivion")
            doc = ET.SubElement(root, "leak")
            for binario in conteudo:
                binario = binario.encode('utf-16', 'replace').decode('utf-16')
                ET.SubElement(doc, f"field", name="leak").text = binario
            tree = ET.ElementTree(root)
            tree.write(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml')

        if xsl_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xml.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xml.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xml',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.xml',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.xml', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

    if db_f == 'PySide2.QtCore.Qt.CheckState.Checked':
        hash_file = random.getrandbits(23)

        if db_ocult == 'PySide2.QtCore.Qt.CheckState.Checked':
            x = 0
            conn = sqlite3.connect(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db')
            c = conn.cursor()

            try:
                c.execute('''CREATE TABLE IF NOT EXISTS leaks
                                          (email text, password text, origin text, data text, source text)''')
                for binario in conteudo:

                    if not 'N/A' in binario:
                        binario = binario.encode('utf-16', 'replace').decode('utf-16')
                        binario_repor = binario.split(':')
                        senha = binario_repor[1]
                        senha_ocult = ocultar_senhas(senha)
                        binario = binario.replace(senha, senha_ocult)
                    binario_lit = binario.split(':')
                    y = (binario_lit[0])
                    c.execute(f'''INSERT INTO leaks
                                             VALUES('{binario_lit[0]}', '{binario_lit[1]}', '{binario_lit[2]}', '{binario_lit[-1]}', '{binario}')''')

                conn.commit()
                conn.close()


            except:
                string_grande(criar=db_f, criar_ocult=db_ocult, criar_cript=db_cript, documentos=documentos,
                              nome=nome, data_atual=data_atual, binario=conteudo, ssh_enviar=ssh_enviar,
                              id_pasta=id_pasta,
                              gdrive=gdrive, aws_s3_gerar=aws_s3_gerar, serverx=serverx, key=key)

        if db_ocult == 'PySide2.QtCore.Qt.CheckState.Unchecked':

            try:
                x = 0
                conn = sqlite3.connect(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db')
                c = conn.cursor()

                c.execute('''CREATE TABLE IF NOT EXISTS leaks
                              (email text, password text, origin text, data text, source text)''')

                for binario in conteudo:
                    binario_lit = binario.split(':')
                    y = (binario_lit[0])
                    c.execute(f'''INSERT INTO leaks
                                 VALUES('{binario_lit[0]}', '{binario_lit[1]}', '{binario_lit[2]}', '{binario_lit[-1]}', '{binario}')''')

                conn.commit()
                conn.close()

            except:
                string_grande(criar=xlsx_f, criar_ocult=xlsx_ocult, criar_cript=xlsx_cript, documentos=documentos,
                              nome=nome, data_atual=data_atual, binario=conteudo, ssh_enviar=ssh_enviar,
                              id_pasta=id_pasta,
                              gdrive=gdrive, aws_s3_gerar=aws_s3_gerar, serverx=serverx, key=key)

        if db_cript == 'PySide2.QtCore.Qt.CheckState.Checked':
            criptografar(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db')
            os.remove(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db')

            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.db.encrypted',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db.encrypted')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db.encrypted',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.db.encrypted', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)

        else:
            if gdrive == 'PySide2.QtCore.Qt.CheckState.Checked':
                subir_arquivo_drive_raiz(id_pasta, f'{nome}resultados_oblivion_{data_atual}_{hash_file}.db',
                                         'text/plain',
                                         f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db')

            if ssh_enviar and serverx:
                ssh_enviar_arquivo(f'{documentos}/Oblivion/{nome}resultados_oblivion_{data_atual}_{hash_file}.db',
                                   f'{nome}resultados_oblivion_{data_atual}_{hash_file}.db', key=key)

            if aws_s3_gerar and serverx:
                subir_bucket_s3(f'{documentos}/Oblivion/', key=key)
